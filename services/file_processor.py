import os
import logging
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from io import BytesIO

logger = logging.getLogger(__name__)

class FileProcessor:
    """Service for processing uploaded files and extracting text content"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
    
    @staticmethod
    def is_allowed_file(filename):
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in FileProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text_from_pdf(file_content):
        """Extract text from PDF file content"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content):
        """Extract text from DOCX file content"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def process_file(file_obj, upload_folder):
        """Process uploaded file and extract text"""
        try:
            # Secure the filename
            filename = secure_filename(file_obj.filename)
            if not filename:
                raise ValueError("Invalid filename")
            
            # Check if file is allowed
            if not FileProcessor.is_allowed_file(filename):
                raise ValueError("File type not allowed")
            
            # Read file content
            file_content = file_obj.read()
            file_obj.seek(0)  # Reset file pointer
            
            # Get file extension
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            # Extract text based on file type
            if file_extension == 'pdf':
                extracted_text = FileProcessor.extract_text_from_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                extracted_text = FileProcessor.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Save file to uploads folder
            file_path = os.path.join(upload_folder, filename)
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            return {
                'filename': filename,
                'file_type': file_extension,
                'extracted_text': extracted_text,
                'file_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            raise Exception(f"File processing failed: {str(e)}")
