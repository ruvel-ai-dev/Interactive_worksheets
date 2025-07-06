#!/usr/bin/env python3
"""
Create a test DOCX file for testing
"""

from docx import Document

def create_test_worksheet():
    doc = Document()
    
    # Add title
    title = doc.add_heading('Simple Addition Practice', 0)
    
    # Add content
    doc.add_paragraph('Learn to add single-digit numbers.')
    doc.add_paragraph('')
    
    # Add examples
    doc.add_paragraph('Examples:')
    doc.add_paragraph('1 + 2 = 3')
    doc.add_paragraph('2 + 3 = 5') 
    doc.add_paragraph('4 + 1 = 5')
    doc.add_paragraph('3 + 3 = 6')
    doc.add_paragraph('')
    
    # Add practice problems
    doc.add_paragraph('Practice Problems:')
    doc.add_paragraph('• What is 2 + 4?')
    doc.add_paragraph('• Calculate 1 + 5')
    doc.add_paragraph('• Find the sum of 3 + 2')
    doc.add_paragraph('• What is 4 + 3?')
    doc.add_paragraph('• Add 1 + 1')
    doc.add_paragraph('')
    
    # Add explanation
    doc.add_paragraph('Addition helps us combine numbers together to find the total amount. When we add two numbers, we are finding how many we have altogether.')
    
    # Save the document
    doc.save('test_worksheet.docx')
    print("✅ Created test_worksheet.docx")

if __name__ == '__main__':
    create_test_worksheet()